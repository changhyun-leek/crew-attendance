from pathlib import Path
from zipfile import ZipFile
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "새벽이슬_출석관리_사용설명서.docx"
FONT = "Malgun Gothic"  # Korean-language named override for compact_reference_guide.
PURPLE = "3B0764"       # Church product brand override.
PRIMARY = "5B21B6"
GOLD = "B7791F"
INK = "21172E"
MUTED = "615A6B"
LIGHT = "F4F1F8"
TABLE_HEAD = "E8EEF5"
GREEN = "087F5B"
RED = "C92A2A"
DXA_WIDTH = 9360
DXA_INDENT = 120


def set_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == DXA_WIDTH
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(DXA_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(DXA_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def style_cell_text(cell, size=9.5, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_font(run, size=size, color=color, bold=bold)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, PRIMARY),
        ("Heading 2", 13, 14, 7, PRIMARY),
        ("Heading 3", 12, 10, 5, PURPLE),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    level.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}[level], color=PRIMARY if level < 3 else PURPLE, bold=True)
    return p


def add_para(doc, text, bold_prefix=None, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_font(a, color=color, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        set_font(b, color=color)
    else:
        r = p.add_run(text)
        set_font(r, color=color)
    return p


def add_callout(doc, title, body, tone="purple"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [DXA_WIDTH])
    cell = table.cell(0, 0)
    mark_header_row(table.rows[0])
    set_cell_shading(cell, "F1EAFE" if tone == "purple" else "FFF8E6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=11, color=PRIMARY if tone == "purple" else GOLD, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_font(r2, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, TABLE_HEAD)
        style_cell_text(cell, size=9.2, bold=True, color=PURPLE, align=(aligns or [WD_ALIGN_PARAGRAPH.LEFT] * len(headers))[index])
    for values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(values):
            cells[index].text = text
            style_cell_text(cells[index], size=9.1, align=(aligns or [WD_ALIGN_PARAGRAPH.LEFT] * len(headers))[index])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def set_header_footer(section):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header.add_run("새벽이슬 청소년부 출석관리")
    set_font(left, size=8.5, color=MUTED, bold=True)
    right = header.add_run("\t사용설명서")
    set_font(right, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left_f = footer.add_run("생명샘동천교회 새벽이슬 청소년부")
    set_font(left_f, size=8, color=MUTED)
    page = footer.add_run("\t")
    set_font(page, size=8, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page._r.extend([fld_begin, instr, fld_end])


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_header_footer(section)
    configure_styles(doc)
    # Word's decimal list markers can disappear with Korean font fallback on some
    # Windows installations. Use the same explicit Unicode bullet definition for
    # all procedural lists so every step remains visibly marked after rendering.
    number_id = add_numbering(doc, bullet=True)
    bullet_id = add_numbering(doc, bullet=True)

    # Editorial cover pattern for a compact_reference_guide manual.
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("생명샘동천교회")
    set_font(r, size=11, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("새벽이슬 청소년부\n출석관리 정식 사용설명서")
    set_font(r, size=29, color=PURPLE, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(48)
    r = sub.add_run("크루교사 · 임원교사 · 보조교사를 위한 운영 안내")
    set_font(r, size=13, color=MUTED)
    add_callout(doc, "이 설명서의 목적", "처음 사용하는 교사도 화면에 적힌 순서대로 출석을 체크하고, 임원은 컴퓨터에서 전체 기록과 보고사항을 안전하게 관리할 수 있도록 안내합니다.")
    add_para(doc, "운영용 · 2026년 7월", color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "운영 주소  https://changhyun-leek.github.io/crew-attendance/", color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    page_break(doc)
    add_heading(doc, "1. 10분 안에 시작하기", 1)
    add_callout(doc, "가장 쉬운 시작", "데스크톱의 새벽이슬 출석관리 바로가기를 더블클릭합니다. 설명서와 관리 파일은 새벽이슬 출석관리 안내 폴더에서 엽니다. 휴대폰에서는 운영 주소를 즐겨찾기 또는 홈 화면에 추가합니다.")
    add_heading(doc, "첫 화면에서 역할 선택", 2)
    add_table(doc, ["역할", "주요 기기", "할 수 있는 일", "제한"], [
        ["크루교사", "휴대폰", "담당 크루 출석, 사유·비고, 학생 관리", "다른 크루 접근 불가"],
        ["임원교사", "컴퓨터", "전체 출석, 공지, 임시 조사, 교사·크루·학생 관리", "임원 PIN 필요"],
        ["보조교사", "휴대폰", "선택 크루의 오늘 출석과 당일 사유", "과거 기록·지속 비고 불가"],
    ], [1440, 1440, 4140, 2340])
    add_heading(doc, "처음 접속하는 순서", 2)
    for text in (
        "크루교사, 임원교사, 보조교사 중 본인의 역할을 누릅니다.",
        "크루교사와 임원교사는 이름을 누르거나 검색창에서 이름을 찾습니다.",
        "숫자 PIN 4~6자리를 입력하고 확인 표시를 누릅니다.",
        "개인 기기에서는 로그인이 최대 30일 유지됩니다. 공용 기기에서는 사용 후 반드시 로그아웃합니다.",
    ):
        add_list_item(doc, text, number_id)
    add_para(
        doc,
        "빠른 찾기  |  크루교사 2장 · 보조교사 3장 · 임원 4~6장 · 보안·오류 7~8장 · 데스크톱·수정 9장",
        bold_prefix="빠른 찾기",
        color=MUTED,
        after=0,
    )

    page_break(doc)
    add_heading(doc, "2. 크루교사 출석체크", 1)
    add_para(doc, "크루교사 화면은 휴대폰에서 한 손으로 사용할 수 있도록 학생 이름과 큰 출석·결석 버튼을 한 줄에 배치했습니다.")
    add_heading(doc, "출석을 시작하는 순서", 2)
    for text in (
        "첫 화면에서 크루교사를 누릅니다.",
        "이름을 직접 누르거나 이름 검색에 이름 일부를 입력합니다.",
        "PIN을 입력한 뒤 화면 상단의 날짜와 크루명을 확인합니다.",
        "학생마다 출석 또는 결석을 누릅니다. 같은 상태를 다시 누르면 미체크로 돌아갑니다.",
        "결석 이유나 연락 결과가 있으면 사유·비고를 누르고 내용을 저장합니다.",
        "임원이 만든 임시 항목이 보이면 학생별 값을 선택합니다.",
        "미체크가 0명인지 확인하고 보고 문구 복사를 사용합니다.",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "저장 완료 확인", "학생별 표시가 저장 중에서 저장됨으로 바뀌었는지 확인합니다. 인터넷 연결 경고 또는 저장 실패가 보이면 저장된 것으로 판단하지 않습니다.", tone="gold")
    add_heading(doc, "학생 추가와 상태 관리", 2)
    for text in (
        "새 학생은 학생 관리에서 이름을 입력해 추가합니다.",
        "오래 나오지 못하는 학생은 장기결석으로 바꾸고, 복귀하면 활동으로 되돌립니다.",
        "더 이상 크루에 속하지 않는 학생은 삭제하지 않고 퇴실로 처리합니다.",
        "학생 상태가 바뀌어도 과거 출석 기록은 그대로 유지됩니다.",
    ):
        add_list_item(doc, text, bullet_id)

    page_break(doc)
    add_heading(doc, "3. 보조교사와 학생 보고사항", 1)
    add_heading(doc, "보조교사 출석체크", 2)
    add_para(doc, "보조교사는 개인 계정 없이 본인 이름을 입력하고 선택한 크루의 오늘 출석만 체크합니다. 화면 상단에는 현재 입력한 보조교사 이름이 표시됩니다.")
    for text in (
        "첫 화면에서 보조교사를 누르고 출석할 크루를 선택합니다.",
        "최근 사용 이름을 고르거나 본인 이름을 두 글자 이상 입력합니다.",
        "오늘 출석·결석, 당일 결석 사유, 연락 상태, 임시 확인 항목을 입력합니다.",
        "이름이 잘못되었으면 화면 상단 이름에서 수정합니다. 이전 이름과 수정 이름은 변경 기록에 남습니다.",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "보조교사 권한", "과거 출석, 지속 학생 비고, 학생 관리, 다른 크루 자료와 내보내기 기능에는 접근할 수 없습니다.")
    add_heading(doc, "사유와 비고를 구분하는 기준", 2)
    add_table(doc, ["입력 위치", "언제 사용하나요?", "예시", "누가 볼 수 있나요?"], [
        ["결석·특수 사유", "해당 날짜에만 필요한 이유", "감기, 가족 일정", "담당교사·임원"],
        ["연락 상태", "학생에게 연락했는지 확인", "연락 안 됨, 연락 완료", "담당교사·임원"],
        ["학생 비고", "앞으로도 계속 알아야 할 사항", "장학금 서류 확인", "담당교사·임원"],
        ["임시 확인 항목", "기간 한정 조사", "수련회 신청/미신청/고려중", "교사 입력, 임원 집계"],
    ], [1650, 2580, 2670, 2460])
    page_break(doc)
    add_heading(doc, "4. 임원교사 전체 현황과 내보내기", 1)
    add_para(doc, "임원 화면은 컴퓨터 사용을 기준으로 왼쪽 메뉴, 상단 필터, 중앙 데이터 표를 배치했습니다. 태블릿과 휴대폰에서는 메뉴가 접히지만 대량 확인과 내보내기는 컴퓨터 사용을 권장합니다.")
    add_heading(doc, "전체 출석을 확인하는 순서", 2)
    for text in (
        "첫 화면에서 임원교사를 누르고 임원 이름과 PIN으로 로그인합니다.",
        "전체 현황에서 운영 크루, 등록 학생, 출석, 결석과 출석률을 확인합니다.",
        "시작일, 종료일, 크루, 학생, 상태, 비고 필터를 필요한 만큼 설정합니다.",
        "조회를 눌러 표를 갱신합니다.",
        "노란색 행과 확인 필요 표식이 있는 학생의 사유·연락·비고를 확인합니다.",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "확인 필요 학생", "결석 사유, 연락 안 됨 또는 학생 비고가 있는 학생은 표에서 노란색으로 표시되어 먼저 확인할 수 있습니다.")
    add_heading(doc, "데이터 내보내기", 2)
    for text in (
        "TXT: 단톡방 또는 일반 문서에 붙이기 쉬운 보고 문구입니다.",
        "CSV: Windows Excel에서 한글이 깨지지 않도록 UTF-8 BOM으로 저장됩니다.",
        "Excel용 표 복사: 버튼을 누른 뒤 Excel 첫 셀에 붙여넣습니다.",
    ):
        add_list_item(doc, text, bullet_id, bold_prefix=text.split(":")[0] + ":")
    add_para(doc, "CSV 열 순서: 출석일, 크루, 학생명, 학생상태, 출석상태, 체크자구분, 체크자명, 체크시각, 최종수정시각", bold_prefix="CSV 열 순서:")

    page_break(doc)
    add_heading(doc, "5. 임원교사 관리 기능", 1)
    add_heading(doc, "크루 관리", 2)
    for text in (
        "크루명과 운영 연도를 입력해 새 크루를 등록합니다.",
        "크루 행에서 담당교사 이름을 선택하고 배정 저장을 누릅니다.",
        "운영 종료는 확인창을 거쳐 처리합니다. 과거 기록은 삭제되지 않습니다.",
    ):
        add_list_item(doc, text, number_id)
    add_heading(doc, "교사 관리", 2)
    for text in (
        "교사 이름, 처음 사용할 PIN, 담당교사 또는 임원교사 역할을 선택해 등록합니다.",
        "PIN을 잊은 교사의 새 PIN을 입력하고 PIN 초기화를 누릅니다.",
        "더 이상 사용하지 않는 계정은 비활성화하고, 필요하면 다시 활성화합니다.",
        "현재 로그인한 임원 본인 계정은 실수로 비활성화할 수 없습니다.",
    ):
        add_list_item(doc, text, number_id)
    add_heading(doc, "학생 관리", 2)
    for text in (
        "학생 이름과 소속 크루를 선택해 등록합니다.",
        "학생 이름 검색으로 학생을 빠르게 찾습니다.",
        "상태를 활동, 장기결석, 퇴실 중 하나로 변경합니다.",
        "다른 크루가 있을 때 이동할 크루를 선택하고 크루 이동을 누릅니다.",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "삭제 대신 상태 변경", "학생과 교사를 직접 삭제하지 않고 상태와 활성 여부를 관리해 과거 출석과 변경 기록을 보존합니다.", tone="gold")

    page_break(doc)
    add_heading(doc, "6. 공지·임시 확인 항목·의견함", 1)
    add_heading(doc, "교사에게 공지 올리기", 2)
    for text in (
        "공지·임시 항목 메뉴를 엽니다.",
        "종류에서 공지를 선택합니다.",
        "제목, 설명, 시작일과 종료일을 입력해 게시합니다.",
    ):
        add_list_item(doc, text, number_id)
    add_heading(doc, "학생별 임시 확인 칸 만들기", 2)
    for text in (
        "종류에서 학생별 임시 확인 칸을 선택합니다.",
        "제목과 교사에게 보일 설명을 입력합니다.",
        "선택지를 쉼표로 구분합니다. 예: 신청, 미신청, 고려중.",
        "필수 여부와 게시 기간을 정해 출석 화면에 게시합니다.",
        "기간이 끝나면 화면에서 자동으로 숨겨지며 입력된 응답은 보존됩니다.",
    ):
        add_list_item(doc, text, number_id)
    add_heading(doc, "오류·개선 의견 처리", 2)
    for text in (
        "교사가 역할, 화면, 문제 상황과 원하는 개선점을 적어 보냅니다.",
        "임원은 의견함에서 새 의견을 확인하고 확인 중으로 변경합니다.",
        "수정 담당자가 데모 화면과 테스트로 변경을 확인합니다.",
        "운영 배포 후 임원이 의견 상태를 완료로 변경합니다.",
    ):
        add_list_item(doc, text, number_id)

    page_break(doc)
    add_heading(doc, "7. 로그인·PIN·개인정보 안전수칙", 1)
    add_table(doc, ["상황", "올바른 사용"], [
        ["개인 휴대폰·컴퓨터", "로그인을 최대 30일 유지해 반복 입력을 줄입니다."],
        ["교회 공용 기기", "업무가 끝나면 왼쪽 또는 상단의 로그아웃을 누릅니다."],
        ["PIN 5회 오류", "10분 동안 잠기므로 반복 입력하지 말고 기다립니다."],
        ["PIN 분실", "임원교사가 교사 관리에서 새 PIN으로 초기화합니다."],
        ["학생 개인정보", "사역에 필요한 최소 범위만 기록하고 화면 캡처에서는 가립니다."],
    ], [2340, 7020])
    add_heading(doc, "기억해야 할 원칙", 2)
    for text in (
        "교사 PIN은 단톡방, 사용설명서, 메모 사진에 남기지 않습니다.",
        "Supabase 서비스 권한 키와 초기 관리자 코드는 GitHub에 올리지 않습니다.",
        "학생은 삭제하지 않고 장기결석 또는 퇴실로 전환합니다.",
        "공용 기기에서는 반드시 로그아웃합니다.",
    ):
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "보안상 문서에 적지 않는 것", "현재 사용 중인 PIN과 비밀키는 이 설명서에 포함하지 않습니다. 초기 PIN 전달과 변경은 임원에게 별도로 확인합니다.", tone="gold")

    page_break(doc)
    add_heading(doc, "8. 문제가 생겼을 때", 1)
    add_table(doc, ["상황", "먼저 할 일", "그다음"], [
        ["인터넷 연결 경고", "출석 내용을 임시 메모", "연결 복구 후 다시 저장"],
        ["저장 실패 표시", "학생과 항목 확인", "재시도 후 의견함에 기록"],
        ["PIN 5회 오류", "10분 기다림", "임원에게 PIN 초기화 요청"],
        ["보조교사 이름 오류", "상단 이름 눌러 수정", "변경 이력에서 확인"],
        ["화면이 예전 모습", "브라우저 새로고침", "컴퓨터는 Ctrl+F5"],
        ["긴급 출석 업무", "임원에게 즉시 알림", "종이·단톡방 기록 후 입력"],
    ], [2520, 3420, 3420])
    add_heading(doc, "오류를 신고할 때", 2)
    for text in (
        "크루교사, 임원교사, 보조교사 중 어떤 역할이었는지",
        "어느 화면에서 어떤 버튼을 어떤 순서로 눌렀는지",
        "기대한 결과와 실제로 보인 결과",
        "오류가 발생한 날짜와 대략적인 시각",
        "가능하면 휴대폰 기종과 개인정보를 가린 화면 캡처",
    ):
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "긴급 상황", "당일 출석 업무를 멈출 수 없으면 임원에게 알리고 종이나 단톡방에 임시 기록한 뒤, 인터넷이 복구되면 시스템에 다시 입력합니다.", tone="gold")

    page_break(doc)
    add_heading(doc, "9. 데스크톱 바로가기와 수정 요청", 1)
    add_para(doc, "데스크톱의 새벽이슬 출석관리 안내 폴더에는 운영 사이트, 이 설명서, 프로그램 관리 폴더, 오류 요청 예시가 함께 있습니다.")
    add_table(doc, ["데스크톱 항목", "용도"], [
        ["새벽이슬 출석관리", "운영 사이트를 기본 웹 브라우저로 엽니다."],
        ["새벽이슬 출석관리 사용설명서", "교사와 임원이 따라 보는 Word 설명서입니다."],
        ["프로그램 관리 폴더", "코드, 백업, 배포 문서와 미리보기를 엽니다."],
        ["오류·개선 요청 작성 예시", "Codex에 전달할 내용을 빠뜨리지 않도록 도와줍니다."],
    ], [3180, 6180])
    add_heading(doc, "수정을 요청하는 가장 쉬운 순서", 2)
    for text in (
        "사이트의 오류·개선 의견 보내기로 현장 의견을 접수합니다.",
        "임원 의견함에서 새 의견을 확인 중으로 변경합니다.",
        "이 Codex 대화에서 새벽이슬 출석관리에서 원하는 내용을 수정해 달라고 요청합니다.",
        "개발_미리보기_실행.bat으로 데모 화면을 확인합니다. 데모 접속 정보는 관리 폴더의 유지보수 가이드에서 확인합니다.",
        "테스트에 통과한 변경만 운영 사이트에 배포합니다.",
        "현장 확인 후 의견 상태를 완료로 변경합니다.",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "운영 데이터 보호", "미리보기의 데모 데이터와 실제 Supabase 운영 데이터를 구분해 검사합니다. 데이터 구조 변경 전에는 반드시 기존 출석을 백업합니다.")

    page_break(doc)
    add_heading(doc, "10. 운영 체크리스트", 1)
    add_heading(doc, "출석체크를 마치기 전", 2)
    for text in (
        "화면 상단의 날짜와 크루명이 올바르다.",
        "모든 학생이 출석 또는 결석으로 체크되었다.",
        "결석 학생의 사유와 연락 상태를 필요한 만큼 입력했다.",
        "임원이 요청한 임시 확인 항목을 빠뜨리지 않았다.",
        "저장 실패 또는 인터넷 연결 경고가 없다.",
        "보고 문구를 복사해 필요한 곳에 보고했다.",
    ):
        add_list_item(doc, text, bullet_id)
    add_heading(doc, "임원이 주기적으로 확인할 것", 2)
    for text in (
        "확인 필요 표시가 있는 학생과 연락 안 됨 학생",
        "학생 비고의 장학금·상담·특별 보고사항",
        "공지와 임시 확인 항목의 종료일",
        "새 오류·개선 의견과 처리 상태",
        "비활성 교사, 종료 크루, 장기결석·퇴실 학생 정리",
        "CSV 내보내기와 기존 출석 건수의 이상 여부",
    ):
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "마지막 확인", "공용 기기라면 로그아웃합니다. 개인 기기에서는 자동 로그인을 이용하되, 다른 사람에게 기기를 빌려줄 때는 먼저 로그아웃합니다.", tone="gold")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "새벽이슬 청소년부 출석관리 사용설명서"
    doc.core_properties.subject = "크루교사, 임원교사, 보조교사 사용 안내"
    doc.core_properties.author = "생명샘동천교회 새벽이슬 청소년부"
    doc.save(OUT)
    audit(OUT)
    print(OUT)


def audit(path):
    from lxml import etree
    with ZipFile(path) as archive:
        xml = etree.fromstring(archive.read("word/document.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        tables = xml.xpath("//w:tbl", namespaces=ns)
        assert tables, "No tables found"
        for table in tables:
            tbl_w = table.xpath("./w:tblPr/w:tblW/@w:w", namespaces=ns)
            tbl_ind = table.xpath("./w:tblPr/w:tblInd/@w:w", namespaces=ns)
            grid = [int(value) for value in table.xpath("./w:tblGrid/w:gridCol/@w:w", namespaces=ns)]
            assert tbl_w == [str(DXA_WIDTH)], tbl_w
            assert tbl_ind == [str(DXA_INDENT)], tbl_ind
            assert sum(grid) == DXA_WIDTH, grid
            for row in table.xpath("./w:tr", namespaces=ns):
                cell_widths = [int(value) for value in row.xpath("./w:tc/w:tcPr/w:tcW/@w:w", namespaces=ns)]
                assert cell_widths == grid, (cell_widths, grid)
        assert xml.xpath("count(//w:numPr)", namespaces=ns) >= 10


if __name__ == "__main__":
    build()
